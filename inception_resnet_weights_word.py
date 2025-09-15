import torch
from facenet_pytorch import InceptionResnetV1
from docx import Document

# Load pre-trained InceptionResnetV1 model
model = InceptionResnetV1(pretrained='vggface2').eval()

# Create a Word document
doc = Document()
doc.add_heading('InceptionResnetV1 Model Weights (VGGFace2)', 0)

# Iterate over model parameters and add first 10 values for readability
for name, param in model.named_parameters():
    doc.add_heading(name, level=1)
    values = param.detach().cpu().numpy().flatten()[:10]  # first 10 weights
    doc.add_paragraph(str(values))

# Save the Word document
doc.save('InceptionResnetV1_Weights.docx')
print('Word file saved as InceptionResnetV1_Weights.docx')